import docx
from docx.shared import Pt
import os
import mimetypes
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table

def process_paragraph(para, image_base_url):
    para_html = ""
    for run in para.runs:
        text = run.text
        if text:
            para_html += text
        
        # Check for images in the run
        if 'graphic' in run._element.xml:
            # Extract rId
            for child in run._element.iter():
                if child.tag.endswith('blip'):
                    rId = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rId and image_base_url:
                        para_html += f'<img src="{image_base_url}/{rId}" style="max-width: 100%; height: auto;" />'
    
    if para_html.strip() or (para.text and para.text.strip()):
         return f"<p>{para_html}</p>"
    return ""

def read_docx(file_path, image_base_url=None):
    try:
        doc = docx.Document(file_path)
        chapters = []
        
        html_content = ""
        
        # Iterate over all content in the body (paragraphs and tables)
        for element in doc.element.body.iterchildren():
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                html_content += process_paragraph(para, image_base_url)
            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                html_content += "<table border='1' style='border-collapse: collapse; width: 100%; margin: 1em 0;'>"
                for row in table.rows:
                    html_content += "<tr>"
                    for cell in row.cells:
                        html_content += "<td style='padding: 8px; border: 1px solid var(--border-color); vertical-align: top;'>"
                        for para in cell.paragraphs:
                            html_content += process_paragraph(para, image_base_url)
                        html_content += "</td>"
                    html_content += "</tr>"
                html_content += "</table>"
                
        chapters.append({
            'id': 'doc-content',
            'content': html_content,
            'href': '#'
        })
                    
        return {
            'title': 'Document', # DOCX metadata is often empty/messy, could try doc.core_properties.title
            'chapters': chapters
        }
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return None

def get_docx_image(file_path, image_id):
    try:
        doc = docx.Document(file_path)
        # image_id is the rId (relationship ID)
        
        if image_id in doc.part.rels:
            rel = doc.part.rels[image_id]
            if "image" in rel.target_ref: # Check if it's an image
                 # rel.target_part.blob contains the image data
                 # We need to guess the content type
                 content_type = mimetypes.guess_type(rel.target_ref)[0]
                 return rel.target_part.blob, content_type
                 
        return None, None
    except Exception as e:
        print(f"Error extracting DOCX image: {e}")
        return None, None

def extract_cover_image(file_path):
    try:
        doc = docx.Document(file_path)
        
        # Iterate through the document body to find the first image in reading order
        # This is more accurate than just taking the first relationship
        for element in doc.element.body.iterchildren():
            # Check paragraphs
            if isinstance(element, CT_P):
                for child in element.iter():
                    if child.tag.endswith('blip'):
                        rId = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if rId and rId in doc.part.rels:
                            rel = doc.part.rels[rId]
                            if "image" in rel.target_ref:
                                content_type = mimetypes.guess_type(rel.target_ref)[0]
                                return rel.target_part.blob, content_type
                                
            # Check tables (images might be in the first table)
            elif isinstance(element, CT_Tbl):
                for child in element.iter():
                    if child.tag.endswith('blip'):
                        rId = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if rId and rId in doc.part.rels:
                            rel = doc.part.rels[rId]
                            if "image" in rel.target_ref:
                                content_type = mimetypes.guess_type(rel.target_ref)[0]
                                return rel.target_part.blob, content_type
                
        return None, None
    except Exception as e:
        print(f"Error extracting DOCX cover: {e}")
        return None, None


